"""Render Markdown-like CR content to a Word (.docx) document."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


def write_markdown_as_docx(markdown: str, path: Path) -> None:
    """Best-effort conversion: headings, bullets, and paragraphs (no raw HTML)."""
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        if line.strip().startswith("```"):
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif re.match(r"^[\-\*]\s+", line):
            text = line.lstrip("-* ").strip()
            try:
                p = doc.add_paragraph(text, style="List Bullet")
            except (KeyError, ValueError):
                p = doc.add_paragraph(f"• {text}")
            p.paragraph_format.space_after = Pt(2)
        elif re.match(r"^\d+\.\s+", line):
            text = re.sub(r"^\d+\.\s+", "", line).strip()
            try:
                p = doc.add_paragraph(text, style="List Number")
            except (KeyError, ValueError):
                p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(2)
        elif line.strip().startswith("|") and "|" in line[1:]:
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.space_after = Pt(2)
        else:
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.space_after = Pt(3)

    doc.save(str(path))
